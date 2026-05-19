import os
import warnings
import pandas as pd


# nettoyer la sortie du terminal: 
warnings.filterwarnings(
    "ignore",
    message="Workbook contains no default style, apply openpyxl's default"
) 
 
##### Imports parsers fichiers excel
from app.parsers.stat_activite import parse_stat_activite_file
from app.parsers.effectifs import parse_effectifs_file
from app.parsers.stats_standard import parse_stats_standard_file
from app.parsers.bilan_actions import parse_bilan_actions_file
from app.parsers.pssm import parse_pssm_file
from app.parsers.psy import parse_psy_file
from app.parsers.infirmiere import parse_ide_file


##### Imports services d'indicateurs
# Chaque fct compute_*_indicators() prend un DataFrame et retourne un dict de métriques (Series pandas, entiers, flottants...)
from app.services.indicator_service import (
    compute_stat_activite_indicators,
    compute_effectifs_indicators,
    compute_stats_standard_indicators,
    compute_bilan_actions_indicators,
    compute_pssm_indicators,
    compute_css_indicators,
    compute_bilans_professionnels_indicators,
    compute_psy_indicators
)

###### Imports graphiques
# Chaque fonction plot_*() génère un graphique matplotlib et le sauvegarde automatiquement dans output/charts/<nom>.png
from app.charts.stat_activite_charts import (
    plot_consultations_par_centre,
    plot_top_motifs,
    plot_repartition_sexe,
)
from app.charts.activite_charts import (
    plot_top_nationalites_hors_france,
)
from app.charts.activite_medicale_charts import ( 
    plot_evolution_activite_medicale,
    plot_repartition_activite_medicale_annee,
    plot_motifs_medecine_generale_charts,
)
from app.charts.amenagements_charts import (
    plot_evolution_amenagements,
    plot_reparition_amenagements,
)
from app.charts.bilan_actions_charts import (
    plot_actions_par_theme,
    plot_consommables_bilan_actions,
    plot_actions_par_site_lisible
)
from app.charts.bilans_charts import (
    plot_bilans_par_composante,
    plot_bilans_internationaux,
    plot_bilans_par_filiere,
)
from app.charts.infirmier_charts import (
    plot_repartition_activite_infirmiere,
    plot_activite_infirmiere_compare,
)
from app.charts.pssm_charts import (
    plot_pssm_sessions,
    plot_pssm_lastest_year,
)
from app.charts.psy_charts import (
    plot_delai_attente_psy,
    plot_problematique_psy,
    plot_duree_suivi,
    plot_consultations_psy_par_composante,
)
from app.charts.psychiatrie import plot_evolution_psychiatrie
from app.charts.consultations_charts import plot_recap_consultations
from app.charts.dietetique_charts import plot_motifs_consultation_dietetique
from app.charts.effectifs_charts import plot_evolution_effectifs
from app.charts.etablissements_charts import plot_top_etablissements
from app.charts.etablissements_conventionnes_charts import plot_etablissements_conventionnes
from app.charts.stat_gyneco import plot_motifs_reels_css, plot_motifs_CSS, plot_prescriptions_css
from app.charts.stats_standard_charts import plot_appels_par_mois
from app.charts.dashboard_pdf import generate_dashboard_pdf
from app.charts.activite_medicale_charts import append_current_year_activite_medicale
from app.charts.amenagements_charts import append_current_year_amenagements
from app.charts.psychiatrie import append_current_year_psychiatrie

##### Imports word_builder et config
from app.services.word_builder import (
    create_document,
    add_title_page,
    add_table_of_contents,
    add_section_heading,
    add_paragraph_text,
    add_bullet_point,
    add_table,
    add_chart_from_file,
    add_page_break,
    add_framed_text,
    add_key_stat,
    save_document,
)

from app.config.intervenants import MEDECINS, INFIRMIERES


charts_dir = "output/charts"

def _chart_path(filename: str) -> str:
    """
    Retourne le chemin complet d'un graphique.
    """
    return os.path.join(charts_dir, filename)


##### Chargement des données (étape 1)
def load_all_data(file_paths: dict) -> dict:
    """
    Charge tous les fichiers Excel et retourne un dictionnaire de DataFrames.
    Sert à valider que les fichiers sont lisibles avant de commencer le rapport
    Si un fichier est absent, la section correspondante sera incomplète mais le reste du rapport sera généré quand meme.
    """
    data = {}

    # stat_activite
    path = file_paths.get("stat_activite")
    if path and os.path.exists(path):
        try:
            data["df_activite"] = parse_stat_activite_file(path)
            print(f"Chargé: {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["df_activite"] = None
    else:
        data["df_activite"] = None
        if path:
            print("Absent: {path}")
    
    # effectifs
    path = file_paths.get("effectifs")
    if path and os.path.exists(path):
        try:
            data["df_effectifs"] = parse_effectifs_file(path)
            print(f"Chargé: {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["df_effectifs"] = None
    else:
        data["df_effectifs"] = None
        if path:
            print("Absent: {path}")
    
    # stats_standard
    path = file_paths.get("stats_standard")
    if path and os.path.exists(path):
        try:
            data["df_stats"] = parse_stats_standard_file(path)
            print(f"Chargé: {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["df_stats"] = None
    else:
        data["df_stats"] = None
        if path:
            print("Absent: {path}")
    
    # bilan_actions
    path = file_paths.get("bilan_actions")
    if path and os.path.exists(path):
        try:
            data["df_actions"] = parse_bilan_actions_file(path)
            print(f"Chargé: {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["df_actions"] = None
    else:
        data["df_actions"] = None
        if path:
            print("Absent: {path}")
    
    # pssm
    path = file_paths.get("pssm")
    if path and os.path.exists(path):
        try:
            data["pssm_sheets"] = parse_pssm_file(path)
            print(f"Chargé: {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["pssm_sheets"] = None
    else:
        data["pssm_sheets"] = None
        if path:
            print("Absent: {path}")
    
    # psy
    path = file_paths.get("psy", "data/raw/stats_psy.xlsx")
    if path and os.path.exists(path):
        try:
            data["df_psy"] = parse_psy_file(path)
            print(f"Chargé: {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["df_psy"] = None
    else:
        data["df_psy"] = None
        if path:
            print(f"Absent: {path}")

    # stat_activite_n_1 (optionnel pour les comparaisons)
    path = file_paths.get("stat_activite_n_1", "data/raw/stat_activite_n_1.xlsx")
    if path and os.path.exists(path):
        try:
            data["df_activite_n_1"] = parse_stat_activite_file(path)
            print(f"Chargé (N-1): {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["df_activite_n_1"] = None
    else:
        data["df_activite_n_1"] = None

    # liste_infirmiere
    path = file_paths.get("liste_infirmiere")
    if path and os.path.exists(path):
        try:
            data["df_liste_infirmiere"] = parse_ide_file(path)
            print(f"Chargé: {path}")
        except Exception as e:
            print(f"Erreur en chargeant {path}: {e}")
            data["df_liste_infirmiere"] = None
    else:
        data["df_liste_infirmiere"] = None
        if path:
             print(f"Absent: {path}")

    return data
    

##### Calcul des indicateurs (étape 2)
def compute_all_indicators(data: dict) -> dict:
    """
    Calcule tous les indicateurs à partir des DataFrames chargés.
    Avoir un seul dict 'stats' à passer au fonctions de construction.
    Gérer proprement les cas où un DataFrame est None.
    """
    stats = {}

    # indicateurs de l'activité principale (toutes consultations confondues)
    if data.get("df_activite") is not None:
        stats["activite"] = compute_stat_activite_indicators(data["df_activite"])
        # on ajoute les stats N-1 si disponibles
        if data.get("df_activite_n_1") is not None:
            stats["activite_n_1"] = compute_stat_activite_indicators(data["df_activite_n_1"])
    else:
        stats["activite"] = {}

    # indicateurs des effectifs étudiants
    if data.get("df_effectifs") is not None:
        stats["effectifs"] = compute_effectifs_indicators(data["df_effectifs"])
    else:
        stats["effectifs"] = {}

    # indicateurs des stats standard (appels téléphoniques, etc)
    if data.get("df_stats") is not None:
        stats["stats_standard"] = compute_stats_standard_indicators(data["df_stats"])
    else:
        stats["stats_standard"] = {}
    
    # indicateurs des bilans d'actions ERS (éducation à la santé)
    if data.get("df_actions") is not None:
        stats["bilan_actions"] = compute_bilan_actions_indicators(data["df_actions"])
    else:
        stats["bilan_actions"] = {}

    
    # indicateurs PSSM (formations premiers secours santé mentale)
    if data.get("pssm_sheets") is not None:
        stats["pssm"] = compute_pssm_indicators(data["pssm_sheets"])
    else:
        stats["pssm"] = {}

    # indicateurs CSS - filtre sur df_activite (motif = "Centre de planification")
    if data.get("df_activite") is not None:
        stats["css"] = compute_css_indicators(data["df_activite"])
    else:
        stats["css"] = {}

    # indicateurs bilans par profession - nécessite les listes de médecins/infirmières
    if data.get("df_activite") is not None:
        stats["bilans_prof"] = compute_bilans_professionnels_indicators(data["df_activite"], MEDECINS, INFIRMIERES)
    else:
        stats["bilans_prof"] = {}
    
    # indicateur psy
    if data.get("df_psy") is not None:
        stats["psy"] = compute_psy_indicators(data["df_psy"])
    else:
        stats["psy"] = {} 

    return stats


##### Génération des graphiques (étape 3)
def generate_all_charts(data: dict, stats: dict):
    """
    Génère tous les graphiques PNG nécessaires pour le rapport dans output/charts.
    S'assure que les graphiques correspondent aux données chargées.
    Si les fichiers Excel ont été mis à jour, les graphique seront mis à jour automatiquement.
    (Si un  graphique échoue, les autres continuent d'etre générés)
    """

    if not os.path.exists(charts_dir):
        os.makedirs(charts_dir)

    # graphiques activite_stats (indicateurs)
    if stats.get("activite"):
        activite_stats = stats["activite"]

        _safe_plot(plot_recap_consultations, activite_stats, "recap_consultations.png")
        _safe_plot(plot_top_nationalites_hors_france, activite_stats, "top_nationalites_hors_france.png")
        # _safe_plot(plot_handicap, activite_stats, "handicap.png")
        # _safe_plot(plot_consultations_par_centre, activite_stats, "consultations_par_centre.png")
        
        # Passage des stats N-1 si présentes pour le graphique top_motifs
        activite_n_1 = stats.get("activite_n_1")
        if activite_n_1:
            _safe_plot(lambda x: plot_top_motifs(x, activite_n_1), activite_stats, "top_motifs.png")
        else:
            _safe_plot(plot_top_motifs, activite_stats, "top_motifs.png")

        _safe_plot(plot_repartition_sexe, activite_stats, "repartition_sexe.png")


    # graphiques df_activite (DataFrame brut)
    if data.get("df_activite") is not None:
        df_activite = data["df_activite"]

        _safe_plot(plot_motifs_medecine_generale_charts,  df_activite,  "motifs_medecine_generale.png")
        # _safe_plot(plot_repartition_activite_depuis_reel, df_activite, "repartition_activite_reelle_IDE.png")
        _safe_plot(plot_duree_suivi, df_activite, "duree_suivi.png")
        _safe_plot(plot_motifs_consultation_dietetique, df_activite, "motifs_consultation_dietetique.png")
        _safe_plot(plot_motifs_CSS, df_activite, "motifs_CSS.png")
        
        delai_psy_path = "data/processed/delai_attente_psy.xlsx"
        if os.path.exists(delai_psy_path):
            _safe_plot(plot_delai_attente_psy, delai_psy_path, "delai_attente_psy.png")

    if data.get("df_psy") is not None:
        _safe_plot(plot_problematique_psy, data["df_psy"], "problematique_psy.png")
        _safe_plot(plot_consultations_psy_par_composante, data["df_psy"], "repartition_psy_composante.png")
        
        # graphiques des bilans de prévention - filtre sur motif = "Bilan de prévention"
        import pandas as pd
        df_bilans = df_activite[df_activite["motif"] == "Bilan de prévention"]
        _safe_plot(plot_bilans_par_composante, df_bilans, "bilans_par_composante.png")
        _safe_plot(plot_bilans_internationaux, df_bilans, "bilans_internationaux.png")
        _safe_plot(plot_bilans_par_filiere, df_bilans, "bilans_par_filiere.png")


    # graphiques df_stats (stats standard)
    if data.get("df_stats") is not None:
        _safe_plot(plot_appels_par_mois, data["df_stats"], "appels_par_mois.png")

    # graphiques df_effectifs
    if data.get("df_effectifs") is not None:
        _safe_plot(plot_evolution_effectifs, data["df_effectifs"], "evolution_effectifs.png")
        _safe_plot(plot_top_etablissements, data["df_effectifs"], "top_etablissements.png")

    # graphiques bilan_actions_stats
    if stats.get("bilan_actions"):
        action_stats = stats["bilan_actions"]
        _safe_plot(plot_actions_par_theme, action_stats, "actions_par_theme.png")
        _safe_plot(plot_consommables_bilan_actions, action_stats, "consommables_bilans_actions.png")
        _safe_plot(plot_actions_par_site_lisible, data["df_actions"], "actions_par_site_lisible.png")

    # graphiques pssm_stats
    if stats.get("pssm"):
        _safe_plot(plot_pssm_sessions, stats["pssm"], "pssm_sessions.png")
        _safe_plot(plot_pssm_lastest_year, data["pssm_sheets"], "pssm_origine_stagiaires.png")
        

    # graphiques css_stats
    if stats.get("css"):
        _safe_plot(plot_motifs_reels_css, stats["css"], "motifs_reels_css.png")
        
        css_prescriptions_path = "data/raw/stats_consultations_css.xlsx"
        if os.path.exists(css_prescriptions_path):
            _safe_plot(plot_prescriptions_css, css_prescriptions_path, "prescriptions_css.png")

    # graohiques prof_stats
    # if stats.get("bilans_prof"):
    #     _safe_plot(plot_bilans_medecins_vs_infirmieres, stats["bilans_prof"], "bilans_medecins_vs_infirmieres.png")

    # Graphiques psychiatrie
    psychiatrie_path = "data/processed/evolution_activite_psychiatrie.xlsx"
    if os.path.exists(psychiatrie_path):
        _safe_plot(plot_evolution_psychiatrie, psychiatrie_path, "evolution_psychiatrie.png")
    
    # Graphiques basés sur des xlsx/Excels manuels dans data/processed
    historique_path = "data/processed/historique_activite_medicale.xlsx"
    if os.path.exists(historique_path):
        _safe_plot(plot_evolution_activite_medicale, historique_path, "evolution_activite_medicale_detail.png")
        _safe_plot(plot_repartition_activite_medicale_annee, historique_path, "repartition_activite_medicale_par_annee.png")

    if data.get("df_liste_infirmiere") is not None:
        _safe_plot(plot_repartition_activite_infirmiere, data["df_liste_infirmiere"], "repartition_activite_infirmiere.png")

    compare_path = "data/processed/activite_infirmiere_compare.xlsx"
    if os.path.exists(compare_path):
        _safe_plot(plot_activite_infirmiere_compare, compare_path, "activite_infirmiere_compare.png")

    
    # Graphiques statiques (données codées dans les modules charts/)
    if data.get("df_effectifs") is not None:
        _safe_plot(plot_etablissements_conventionnes, data["df_effectifs"], "etablissements_conventionnes.png")
    if data.get("df_activite") is not None:
        _safe_plot(plot_reparition_amenagements, data["df_activite"], "repartition_amenagements.png")
    
    amenagements_path = "data/processed/evolutions_amenagements.xlsx"
    if os.path.exists(amenagements_path):
        _safe_plot(plot_evolution_amenagements, amenagements_path, "evolution_amenagements.png")


def _safe_plot(plot_fn, arg, chart_name: str):
    """Permet de continuer la génnération meme si un graphique échoue (capture les exceptions)"""
    try:
        plot_fn(arg)
        print(f"OK : {chart_name}")
    except Exception as e:
        print(f"ERREUR : {chart_name} - {e}")


def _safe_plot_no_args(plot_fn, chart_name: str):
    """Version pour les graphiques sans arguments"""
    try:
        plot_fn()
        print(f"OK : {chart_name}")
    except Exception as e:
        print(f"ERREUR : {chart_name} - {e}")


##### Construction des sections du rapport
def build_synthese_generale(doc, stats: dict):
    """Construit la section "Synthèse générale" du rapport"""

    add_section_heading(doc, "1. Introduction", level=1)
    add_section_heading(doc, "1.1 Synthèse générale", level=2)

    add_paragraph_text(doc, "Ce rapport présente le bilan d'activité du Service de Santé Universitaire (SSU) "
        "pour l'année universitaire 2024-2025. Il couvre l'ensemble des services proposés "
        "aux étudiants : médecine générale, service infirmier, psychologie, psychiatrie, "
        "éducation à la santé, formations PSSM, centre de santé sexuelle et diététique.")
    
    add_paragraph_text(doc, "Les données présentées sont issues des fichiers de suivi d'activité renseignés "
        "tout au long de l'année universitaire. Les graphiques sont générés automatiquement "
        "à partir de ces données.")

    # indicateurs clés
    activite = stats.get("activite", {})
    total = activite.get("total_consultations", "N/A")
    uniques = activite.get("etudiants_uniques", "N/A")
    age_moy = activite.get("age_moyen", None)

    add_key_stat(doc, "Total consultations", f"{int(total)}")
    add_key_stat(doc, "Étudiants uniques", f"{int(uniques)}")
    add_key_stat(doc, "Âge moyen des consultants", f"{age_moy:.1f} ans" if age_moy else "N/A")

    # graphiques de synthèses
    add_chart_from_file(doc, _chart_path("recap_consultations.png"),"")
    add_chart_from_file(doc, _chart_path("top_motifs.png"), "")
    add_chart_from_file(doc, _chart_path("repartition_sexe.png"), "")
    # add_chart_from_file(doc, _chart_path("consultations_par_centre.png"), "Figure 4 - Répartition des consultations par centre")
    # A. Évolution des effectifs
    add_section_heading(doc, "1.2. Évolution des effectifs étudiants", level=2)

    add_chart_from_file(
        doc,
        _chart_path("evolution_effectifs.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("etablissements_conventionnes.png"),
        caption=""
    )

    add_section_heading(doc, "1.3 Évolution de l'activité médicale", level=2)

    add_chart_from_file(
        doc,
        _chart_path("evolution_activite_medicale_detail.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("repartition_activite_medicale.png"),
        caption=""
    )



def build_medecine_generale(doc, stats: dict, data: dict):
    """Construit la section "Médecine générale" du rapport"""
    add_page_break(doc) # chaque section commence sur une nouvelle page
    add_section_heading(doc, "2. Médecine générale", level=1)

    add_paragraph_text(doc, "Le service de médecine générale du SSU assure les consultations courantes, "
        "les bilans de prévention obligatoires et le suivi des étudiants en situation "
        "de handicap ou nécessitant un aménagement d'études.")

    # indicateurs clés
    activite = stats.get("activite", {})
    consult_med = activite.get("consultations_medecine_generale", "N/A")
    consult_bilan = activite.get("consultations_bilans", "N/A")

    add_key_stat(doc, "Consultations médecine générale", f"{str(consult_med)}")
    add_key_stat(doc, "Bilans de prévention", f"{str(consult_bilan)}")

    # 2.1 Motifs de consultation
    add_section_heading(doc, "2.1 Motifs de consultation", level=2)

    add_chart_from_file(
        doc,
        _chart_path("motifs_medecine_generale.png"),
        caption=""
    )

    # 2.2 Bilans de prévention
    add_section_heading(doc, "2.2 Bilans de prévention", level=2)

    add_paragraph_text(
        doc,
        "Les bilans de prévention sont réalisés par les médecins et les infirmières "
        "du SSU. Ils comprennent un bilan de santé général, la mise à jour des "
        "vaccinations et un dépistage des facteurs de risque."
    )

    # indicateurs clés
    bilans_prof  = stats.get("bilans_prof", {})
    bilans_med   = bilans_prof.get("bilans_medecins",    "N/A")
    bilans_inf   = bilans_prof.get("bilans_infirmieres", "N/A")

    add_key_stat(doc, "Bilans réalisés par les médecins",     str(bilans_med))
    add_key_stat(doc, "Bilans réalisés par les infirmières",  str(bilans_inf))

    add_chart_from_file(
        doc,
        _chart_path("bilans_par_composante.png"),
        caption=""
    )

    # add_chart_from_file(
    #     doc,
    #     _chart_path("bilans_medecins_vs_infirmieres.png"),
    #     caption="Figure – Bilans réalisés par les médecins vs les infirmières"
    # )
    
    # 2.3 Profil des étudiants
    add_section_heading(doc, "2.3 Profil des étudiants", level=2)
    add_chart_from_file(
        doc,
        _chart_path("top_nationalites_hors_france.png"),
        caption=""
    )

    # 2.4 Aménagements d'études
    add_section_heading(doc, "2.4 Aménagements d'études", level=2)
    
    add_paragraph_text(
        doc,
        "Le SSU instruit les dossiers d'aménagement d'études pour les étudiants "
        "en situation de handicap ou souffrant d'une pathologie chronique impactant "
        "leur parcours universitaire."
    )

    add_chart_from_file(
        doc,
        _chart_path("evolution_amenagements.png"),
        caption=""
    )
    
    add_chart_from_file(
        doc,
        _chart_path("repartition_amenagements.png"),
        caption=""
    )

def build_service_infirmier(doc, stats:dict, data:dict):
    """Construit la section "Service infirmier" du rapport"""
    add_page_break(doc)
    add_section_heading(doc, "3. Service infirmier", level=1)

    add_paragraph_text(
        doc,
        "Le service infirmier assure les consultations infirmières courantes, les soins, "
        "les vaccinations et les primo-accueils. Les infirmières jouent un rôle clé dans "
        "l'orientation des étudiants vers les services adaptés à leurs besoins."
    )

    # indicateurs clés
    activite = stats.get("activite", {})
    consult_ide = activite.get("consultations_ide", "N/A")
    add_key_stat(doc, "Consultations infirmières (IDE)", str(consult_ide))

    # graphiques
    add_chart_from_file(
        doc,
        _chart_path("repartition_activite_infirmiere.png"),
        caption=""
    )

    # add_chart_from_file(
    #     doc,
    #     _chart_path("repartition_activite_reelle_IDE.png"),
    #     caption="Figure - Répartition de l'activité infirmière (motifs réels déclarés)"
    # )
    
    add_chart_from_file(
        doc,
        _chart_path("activite_infirmiere_compare.png"),
        caption=""
    )

    
def build_psychologie(doc, stats: dict, data: dict):
    """Construit la section "Psychologie" du rapport"""
    add_page_break(doc)
    add_section_heading(doc, "4. Psychologie", level=1)

    add_paragraph_text(
        doc, 
        "Le service de psychologie propose des consultations individuelles pour les "
        "étudiants confrontés à des difficultés psychologiques, émotionnelles ou "
        "relationnelles. L'équipe de psychologues intervient sur plusieurs campus."
    )

    # indicateurs clés
    activite = stats.get("activite", {})
    consult_psy = activite.get("consultations_psychologie", "N/A")
    consult_psy2 = activite.get("consultations_psychiatrie", "N/A")

    add_chart_from_file(
        doc,
        _chart_path("problematique_psy.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("repartition_psy_composante.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("duree_suivi.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("delai_attente_psy.png"),
        caption=""
    )

    # 4.1 Dispositif Santé Psy Etudiants (DSPE)
    add_section_heading(doc, "4.1 Dispositif Santé Psy Etudiants (DSPE)", level=2)

    add_paragraph_text(
        doc,
        "Le Dispositif Santé Psy Étudiants (DSPE) permet aux étudiants de bénéficier "
        "de séances de soutien psychologique remboursées. Le SSU coordonne ce dispositif "
        "en lien avec les psychologues partenaires agréés."
    )

    # cadre encadré pour mettre en valeur le dispositif
    add_framed_text(
        doc, 
        "Le DSPE garantit l'accès à des soins psychologiques pour tous les étudiants, "
        "indépendamment de leur situation financière. Les séances sont entièrement "
        "prises en charge par l'Assurance Maladie."
    )


    # 4.2 Psychiatrie
    add_section_heading(doc, "4.2 Psychiatrie", level=2)

    add_paragraph_text(
        doc,
        "Le service de psychiatrie assure le suivi des étudiants nécessitant une prise "
        "en charge psychiatrique. Il travaille en étroite collaboration avec le service "
        "de psychologie pour faciliter les orientations et garantir la continuité des soins."
    )

    add_chart_from_file(
        doc,
        _chart_path("evolution_psychiatrie.png"),
        caption=""
    )
    



def build_education_sante(doc, stats: dict, data: dict):
    """Construit la section "Éducation à la santé (ERS)" du rapport"""
    add_page_break(doc)
    add_section_heading(doc, "5. Éducation à la santé (ERS)", level=1)

    add_paragraph_text(
        doc,
        "Le pôle Éducation et Promotion de la Santé (ERS) développe des actions de "
        "prévention et d'éducation à la santé à destination des étudiants sur les "
        "différents campus. Ces actions sont menées en partenariat avec les BDE, "
        "les associations étudiantes et les composantes universitaires."
    )

    # indicateurs clés
    bilan         = stats.get("bilan_actions", {})
    nb_actions    = bilan.get("nombre_actions",        "N/A")
    tot_etudiants = bilan.get("total_etudiants_touches", "N/A")

    if isinstance(tot_etudiants, (float, int)) and not isinstance(tot_etudiants, bool):
        tot_etudiants = int(tot_etudiants) # conversion float -> int pour l'affichage si c un nombre

    add_key_stat(doc, "Nombre d'actions menées", str(nb_actions))
    add_key_stat(doc, "Nombre d'étudiants touchés", str(tot_etudiants))

    # 5.1 Répartition par thème
    add_section_heading(doc, "5.1 Répartition par thème", level=2)

    add_chart_from_file(
        doc,
        _chart_path("actions_par_theme.png"),
        caption=""
    )

    # 5.2 Actions par site
    add_section_heading(doc, "5.2 Actions par site universitaire", level=2)

    add_chart_from_file(
        doc,
        _chart_path("actions_par_site_lisible.png"),
        caption=""
    )
    
    # 5.3 consommables distribués
    add_section_heading(doc, "5.3 Consommables et matériel distribués", level=2)

    add_paragraph_text(
        doc,
        "Lors des actions de prévention, le SSU distribue divers supports d'information "
        "et consommables de santé aux étudiants (préservatifs, dépliants, kits prévention...)."
    )
    
    add_chart_from_file(
        doc,
        _chart_path("consommables_bilans_actions.png"),
        caption=""
    )
    

def build_pssm(doc, stats: dict, data: dict):
    """
    Construit la section "Formations PSSM" du rapport.
    """
    add_page_break(doc)
    add_section_heading(doc, "6. Formations PSSM", level=1)

    add_paragraph_text(
        doc,
        "Les formations aux Premiers Secours en Santé Mentale (PSSM) sont organisées "
        "par le SSU pour sensibiliser étudiants et personnels à la santé mentale. "
        "Ces formations certifiantes permettent d'apprendre à reconnaître les signes "
        "d'une détresse psychologique et à intervenir de manière adaptée."
    )

    add_paragraph_text(
        doc,
        "Le programme PSSM est reconnu internationalement. Il s'adresse à toute personne "
        "souhaitant devenir un premier maillon de la chaîne de soutien en santé mentale "
        "dans son environnement (campus, famille, travail...)."
    )

    # indicateurs clés
    pssm           = stats.get("pssm", {})
    nb_sessions    = pssm.get("nombre_sessions",             "N/A")
    nb_participants = pssm.get("total_participants_declares", "N/A")

    if isinstance(nb_participants, float):
        nb_participants = int(nb_participants)

    add_key_stat(doc, "Sessions organisées",  str(nb_sessions))
    add_key_stat(doc, "Participants formés",  str(nb_participants))

    # graphiques
    add_chart_from_file(
        doc,
        _chart_path("pssm_sessions.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("pssm_origine_stagiaires.png"),
        caption=""
    )


def build_centre_sante_sexuelle(doc, stats: dict, data: dict):
    """
    Construit la section "Centre de santé sexuelle (CSS)" du rapport.
    """
    add_page_break(doc)
    add_section_heading(doc, "7. Centre de santé sexuelle", level=1)

    add_paragraph_text(
        doc,
        "Le Centre de Santé Sexuelle (CSS) du SSU propose des consultations de "
        "contraception, de dépistage des infections sexuellement transmissibles (IST) "
        "et de suivi gynécologique dans un cadre bienveillant et confidentiel."
    )

    add_paragraph_text(
        doc,
        "Le centre est ouvert à tous les étudiants sans rendez-vous sur certaines "
        "plages horaires. L'équipe est formée pour accueillir toutes les situations "
        "sans jugement."
    )

    # indicateurs clés
    css       = stats.get("css", {})
    total_css = css.get("total_consultations_css", "N/A")
    add_key_stat(doc, "Consultations Centre de santé sexuelle", str(total_css))

    # 7.1 Motifs de consultation
    add_section_heading(doc, "7.1 Motifs de consultation", level=2)

    add_paragraph_text(
        doc,
        "Les motifs réels de consultation reflètent les besoins concrets des "
        "étudiants fréquentant le CSS."
    )

    add_chart_from_file(
        doc,
        _chart_path("motifs_CSS.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("motifs_reels_css.png"),
        caption=""
    )

    add_chart_from_file(
        doc,
        _chart_path("prescriptions_css.png"),
        caption=""
    )


def build_dietetique(doc, stats: dict, data: dict):
    """
    Construit la section "Diététique et Nutrition" du rapport.
    """
    add_page_break(doc)
    add_section_heading(doc, "8. Diététique et Nutrition", level=1)

    add_paragraph_text(
        doc,
        "Le service de diététique propose des consultations individuelles pour les "
        "étudiants souhaitant améliorer leur alimentation ou souffrant de troubles "
        "du comportement alimentaire. Les consultations sont réalisées par une "
        "diététicienne diplômée d'État."
    )

    # 8.1 Consultations de diététique
    add_section_heading(doc, "8.1 Consultations de diététique", level=2)

    add_bullet_point(doc, "Consultation initiale : reprise de l'historique nutritionnel et enquête alimentaire")
    add_bullet_point(doc, "Bilan des habitudes alimentaires et définition des objectifs")
    add_bullet_point(doc, "Mise en place d'un suivi personnalisé adapté à la demande")
    add_bullet_point(doc, "Orientation vers des spécialistes si nécessaire", level=1)
    add_bullet_point(doc, "Gastro-entérologue, endocrinologue, SUAPS", level=1)
    add_bullet_point(doc, "Diététicienne libérale de ville", level=1)

    add_chart_from_file(
        doc,
        _chart_path("motifs_consultation_dietetique.png"),
        caption=""
    )

    # 8.2 Consultations de nutrition
    add_section_heading(doc, "8.2 Consultations de nutrition", level=2)

    add_paragraph_text(
        doc,
        "Les consultations de nutrition s'adressent aux étudiants avec des besoins "
        "nutritionnels spécifiques : végétarisme, troubles du transit, prise ou perte "
        "de poids, rééquilibrage alimentaire, troubles du comportement alimentaire."
    )

    add_bullet_point(doc, "Surcharge pondérale ou obésité")
    add_bullet_point(doc, "Végétarisme et végétalisme")
    add_bullet_point(doc, "Troubles du comportement alimentaire (TCA)")
    add_bullet_point(doc, "Troubles du transit")
    add_bullet_point(doc, "Rééquilibrage alimentaire et hygiène de vie")
    add_bullet_point(doc, "Demande de prise de poids")


def build_annexes(doc, stats: dict, data: dict):
    """Construit la section "Annexes" du rapport."""
    add_page_break(doc)
    add_section_heading(doc, "Annexes", level=1)


    # A. Appels téléphoniques
    add_section_heading(doc, "A. Appels téléphoniques", level=2)

    add_paragraph_text(
        doc,
        "Le SSU reçoit de nombreux appels téléphoniques au cours de l'année "
        "universitaire. Ces appels concernent la prise de rendez-vous, les "
        "demandes d'information et les situations d'urgence."
    )

    add_chart_from_file(
        doc,
        _chart_path("appels_par_mois.png"),
        caption=""
    )



    

##### POINT D'ENTREE PRINCIPAL
def build_report(file_paths: dict, output_path: str = "output/rapport_ssu.docx", regenerate_charts: bool = True, progress_callback=None) -> str: # retourne le chemin du fichier généré
    """Fonction principale: orchestre la génération complète du rapport Word SSU."""
    """C'est le point d'entrée appelé par l'interface Tkinter."""

    def _progress(step: str, pct: int):
        """Appelle le callback si fourni, sinon affiche dans le terminal."""
        print(f"[{pct:3d}%] {step}")
        if progress_callback:
            progress_callback(step, pct)

    
    # chargement des données 
    _progress("Chargement des fichiers Excel...", 5)
    data = load_all_data(file_paths)

    # calcul des indicateurs
    _progress("Calcul des indicateurs...", 20)
    stats = compute_all_indicators(data)

    # Détection de l'année universitaire en cours (dynamique)
    current_year = "2025-2026" # par défaut
    if data.get("df_activite") is not None:
        df_act = data["df_activite"]
        if "date consultation" in df_act.columns:
            dates = pd.to_datetime(df_act["date consultation"], errors="coerce").dropna()
            if not dates.empty:
                def extract_academic_year(d):
                    return f"{d.year}-{d.year+1}" if d.month >= 8 else f"{d.year-1}-{d.year}"
                academic_years = dates.apply(extract_academic_year)
                if not academic_years.empty:
                    current_year = academic_years.mode()[0]
    
    # Mise à jour des fichiers historiques (uniquement si régénération activée)
    if regenerate_charts:
        _progress("Mise à jour des historiques...", 25)
        if data.get("df_activite") is not None:
            df_act = data["df_activite"]
            try:
                append_current_year_activite_medicale(df_act, "data/processed/historique_activite_medicale.xlsx", current_year)
                append_current_year_amenagements(df_act, "data/processed/evolutions_amenagements.xlsx", current_year)
                append_current_year_psychiatrie(df_act, "data/processed/evolution_activite_psychiatrie.xlsx", current_year)
            except Exception as e:
                print(f"Erreur lors de la mise à jour des historiques: {e}")

    # génération des graphiques
    if regenerate_charts:
        _progress("Génération des graphiques...", 35)
        try:
            generate_all_charts(data, stats)
        except Exception as e:
            print(f"Erreur lors de la génération des graphiques: {e}")
            print("-> Le rapport sera généré avec les graphiques existants.")

    # création du document word
    _progress("Création du document Word...", 50)
    
    # page de titre personnalisée ou par défaut
    # get_resource_path() gère le cas PyInstaller (sys._MEIPASS) et le mode normal
    from app.ui.utils import get_resource_path
    TEMPLATE_PATH = get_resource_path("app/reports/template_page_de_garde.docx")
    doc = create_document(TEMPLATE_PATH)

    if os.path.exists(TEMPLATE_PATH):
        _progress("Utilisation de la page de garde personnalisée...", 52)
        # On ajoute une nouvelle section pour que le reste du document (sommaire, etc.)
        # commence sur une nouvelle page et retrouve les marges normales,
        # sans modifier la mise en page de la page de garde.
        from docx.enum.section import WD_SECTION
        from app.services.word_builder import margin
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.top_margin = margin
        new_section.bottom_margin = margin
        new_section.left_margin = margin
        new_section.right_margin = margin

        # Détacher l'en-tête et le pied de page du template pour qu'ils ne se 
        # répètent pas sur les pages suivantes (au cas où le template en contient)
        from docx.oxml import OxmlElement
        
        for hdr in [new_section.header, new_section.first_page_header, new_section.even_page_header]:
            hdr.is_linked_to_previous = False
            hdr._element.clear()
            hdr._element.append(OxmlElement('w:p'))
            
        for ftr in [new_section.footer, new_section.first_page_footer, new_section.even_page_footer]:
            ftr.is_linked_to_previous = False
            ftr._element.clear()
            ftr._element.append(OxmlElement('w:p'))
        
        # Désactiver les en-têtes différents sur la première page pour cette section (si activé dans le modèle)
        new_section.different_first_page_header_footer = False
    else:
        LOGO_SSU_PATH = "app/reports/logo_ssu.png"
        LOGO_UA_PATH = "app/reports/logo_ua.png"

        _progress("Ajout de la page de titre générée...", 52)
        add_title_page(
            doc, 
            title="Rapport d'Activité du SSU", 
            subtitle="Service de Santé Universitaire - Université d'Angers", 
            year="Année universitaire 2025-2026",
            logo_ssu_path=LOGO_SSU_PATH,
            logo_ua_path=LOGO_UA_PATH
        )


    # table des matières
    _progress("Ajout de la table des matières...", 55)
    sections_toc = [
        {"num": "1", "title": "Synthèse générale", "page": "3"},
        {"num": "2", "title": "Médecine générale", "page": "5"},
        {"num": "3", "title": "Service infirmier", "page": "7"},
        {"num": "4", "title": "Psychologie", "page": "9"},
        {"num": "5", "title": "Éducation à la santé (ERS)", "page": "11"},
        {"num": "6", "title": "Formations PSSM", "page": "13"},
        {"num": "7", "title": "Centre de santé sexuelle", "page": "14"},
        {"num": "8", "title": "Diététique et Nutrition", "page": "16"},
        {"num": "",  "title": "Annexes", "page": "18"},
    ]
    add_table_of_contents(doc, sections_toc)

    # construction des sections
    _progress("Section Synthèse générale...", 60)
    build_synthese_generale(doc, stats)

    _progress("Section Médecine générale...", 65)
    build_medecine_generale(doc, stats, data)

    _progress("Section Service infirmier...", 70)
    build_service_infirmier(doc, stats, data)

    _progress("Section Psychologie...", 75)
    build_psychologie(doc, stats, data)

    _progress("Section Éducation à la santé (ERS)...", 80)
    build_education_sante(doc, stats, data)

    _progress("Section Formations PSSM...", 84)
    build_pssm(doc, stats, data)

    _progress("Section Centre de santé sexuelle...", 87)
    build_centre_sante_sexuelle(doc, stats, data)

    _progress("Section Diététique et Nutrition...", 90)
    build_dietetique(doc, stats, data)

    _progress("Annexes...", 93)
    build_annexes(doc, stats, data)

    # sauvegarde du document
    _progress("Sauvegarde du document word...", 96)
    save_document(doc, output_path)

    # Génération du Dashboard PDF
    _progress("Génération du tableau de bord PDF...", 98)
    try:
        dashboard_data = {
            "demographie": {
                "effectifs": stats.get("effectifs"),
                "etablissements": data["df_effectifs"].to_dict(orient='records') if data.get("df_effectifs") is not None and not data["df_effectifs"].empty else []
            },
            "activite": stats.get("activite"),
            "bilans_prevention": stats.get("bilans_prof"),
            "sante_mentale": {
                "pssm": stats.get("pssm")
            },
            "sante_sexuelle": stats.get("css"),
            "prevention": stats.get("bilan_actions"),
            "stats_standard": stats.get("stats_standard")
        }
        
        output_dir = os.path.dirname(output_path) or "output"
        dashboard_pdf_path = os.path.join(output_dir, "tableau_de_bord_ssu.pdf")
        
        generate_dashboard_pdf(
            data=dashboard_data,
            charts_dir=charts_dir,
            output_path=dashboard_pdf_path,
            year_label=current_year
        )
    except Exception as e:
        print(f"Erreur lors de la génération du Dashboard PDF: {e}")

    abs_path = os.path.abspath(output_path)
    _progress(f"Rapport généré: {abs_path}", 100)

    return abs_path
    

    
    
    
