from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

##### Constantes de couleurs et de style
cyan_primary = RGBColor(0, 169, 206)
cyan_hex = "00A9CE" # meme couleur pour les manip XML python-docx
grey_light = RGBColor(232, 232, 232)
grey_hex = "E8E8E8" # meme couleur pour les manip XML python-docx
white = RGBColor(255, 255, 255)
white_hex = "FFFFFF" # meme couleur pour les manip XML python-docx
black = RGBColor(0, 0, 0)
black_hex = "000000" # meme couleur pour les manip XML python-docx

# Police disponible dans Word
font_family = "Calibri"
# Marge de 1.25 inch
margin = Inches(1.25)

##### Fonctions interne: couleur de fond d'une cellule
def _set_cell_background(cell, hex_color):
    """
    Remplit l'arrière-plan d'une cellule de tableau avec une couleur hexadécimale.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd") # crée l'elt de remplissage OpenXML
    shd.set(qn("w:val"), "clear") # un remplissage uni 
    shd.set(qn("w:color"), "auto") # couleur du motif automatique (pas de motif ici)
    shd.set(qn("w:fill"), hex_color) # couleur de fond réelle
    tcPr.append(shd) # ajoute le remplissage aux propriétés de la cellule
    

##### Création et configuration du document
def create_document(template_path: str = None) -> Document: # Document : objet python-docx
    """
    Crée et configure un nouveau document Word.
    """
    if template_path and os.path.exists(template_path):
        doc = Document(template_path) 
    else:
        doc = Document()
        # appliquer les marges sur toutes les sections du document
        # (un document simple n'a qu'une seule section)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin
    return doc


##### Page de titre
def add_title_page(doc: Document, title: str, subtitle: str="", year: str="", logo_ssu_path: str = None, logo_ua_path: str = None):
    """
    Ajoute une page de titre avec logos en haut et titre au centre.
    """
    # LOGOS ua et ssu en haut
    if logo_ssu_path and os.path.exists(logo_ssu_path):
        p_logo_ssu = doc.add_paragraph()
        p_logo_ssu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo_ssu.paragraph_format.space_after = Pt(6)
        
        run_logo = p_logo_ssu.add_run()
        run_logo.add_picture(logo_ssu_path, width=Inches(1.5))
    
    if logo_ua_path and os.path.exists(logo_ua_path):
        p_logo_ua = doc.add_paragraph()
        p_logo_ua.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo_ua.paragraph_format.space_after = Pt(24)  
        
        run_logo_ua = p_logo_ua.add_run()
        run_logo_ua.add_picture(logo_ua_path, width=Inches(1.5))  # ← PETIT (1.5 inch)

    # titre principal (centré, après les logos)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(24)

    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = cyan_primary
    run.font.name = font_family

    # sous-titre
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(12)

        run_sub = p_sub.add_run(subtitle)
        run_sub.italic = True
        run_sub.font.size = Pt(20)
        run_sub.font.color.rgb = RGBColor(100, 100, 100)
        run_sub.font.name = font_family

    # année
    if year:
        p_year = doc.add_paragraph()
        p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_year.paragraph_format.space_after = Pt(48)

        run_year = p_year.add_run(year)
        run_year.font.size = Pt(12)
        run_year.font.name = font_family
        run_year.font.color.rgb = RGBColor(80, 80, 80)

    # sauts de page
    doc.add_page_break() # séparer la page de titre du reste du document

##### Table des matières
def add_table_of_contents(doc: Document, sections: list):
    """
    Ajoute une table des matières textuelle au document.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    
    # titre de la section
    add_section_heading(doc, "Sommaire", level=1)

    # liste des sections
    for item in sections:
        # support de 2 formats: dict ou string
        if isinstance(item, dict): # cas d'un dict comme {"num": "1", "title": "Médecine générale", "page": "3"}
            num = item.get("num", "")
            title = item.get("title", "")
            page = item.get("page", "")
            line = f"{num}. {title}" if num else title
        else:
            line = str(item)
            page = ""

        p = doc.add_paragraph()
        
        # Ajout d'une tabulation avec des points de suite alignée à droite
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # Texte principal
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = font_family
        
        # Numéro de page (si fourni) après une tabulation
        if page:
            run_page = p.add_run(f"\t{page}")
            run_page.font.size = Pt(11)
            run_page.font.name = font_family
            run_page.font.bold = True
            
        p.paragraph_format.space_after = Pt(14)

    doc.add_page_break() # séparer la table des matières du reste du document        

##### Titres de section
def add_section_heading(doc: Document, title: str, level: int = 1):
    """
    Ajoute un titre de section avec un style graphique cohérent:
        level=1 → titre principal de section   (18pt, gras)
        level=2 → sous-titre                   (14pt, gras)
        level=3 → sous-sous-titre              (12pt, gras, italique)
    """

    # taille de police selon le niveau
    sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12)}
    font_size = sizes.get(level, Pt(12))

    p = doc.add_paragraph()
    run = p.add_run(title)

    run.bold = True
    run.font.size = font_size
    run.font.color.rgb = cyan_primary
    run.font.name = font_family

    if level == 3:
        run.italic = True

    # espacement
    p.paragraph_format.space_before = Pt(14) if level == 1 else Pt(10)
    p.paragraph_format.space_after = Pt(6)

    
    
##### Texte normal
def add_paragraph_text(doc: Document, text: str, bold: bool = False, italic: bool = False):
    """
    Ajoute un paragraphe de texte standard justifié, en noir.
    """

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = font_family
    run.font.color.rgb = black
    
    p.paragraph_format.space_after = Pt(6)


##### Listes
def add_bullet_point(doc: Document, text:str, level: int = 0):
    """
    Ajoute un élément de liste à puce avec un retrait progressif.
    """
    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
        text = "• " + text
        p.paragraph_format.left_indent = Inches(0.25)

    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = font_family
    run.font.color.rgb = cyan_primary

    if level > 0: # indentation supp pour chaque niveau de sous-puce
        p.paragraph_format.left_indent = Inches(0.5 * (level+1))
    
    p.paragraph_format.space_after = Pt(3)
        

##### Tableaux
def add_table(doc: Document, headers: list, rows: list, col_widths: list = None):
    """
    Ajoute un tableau avec en-têtes et lignes de données.
        - Ligne d'en-tete: fond cyan, texte blanc
        - Lignes paires: fond blanc
        - Lignes impaires: fond gris clair

    Returns Table: objet tableau python-docx
    """

    n_cols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=n_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass # Le style n'existe pas dans le template, utiliser le style par défaut

    # ligne d'en-tete
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""

        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = font_family
        run.font.color.rgb = white

        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set_cell_background(cell, cyan_hex)

    # lignes de données
    for row_ind, row_data in enumerate(rows):
        table_row = table.rows[row_ind + 1] # car ligne 0 = en-tete

        # alterner les couleurs de fond
        bg_color = grey_hex if row_ind % 2 == 1 else white_hex
    
        for col_ind, cell_val in enumerate(row_data):
            cell = table_row.cells[col_ind]
            cell.text = str(cell_val) if cell_val is not None else ""

            runs = cell.paragraphs[0].runs
            if runs:
                runs[0].font.size = Pt(10)
                runs[0].font.name = font_family
                
                _set_cell_background(cell, bg_color)

    # largeurs de colonnes
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width
    
    # espace après le tableau
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    
    return table    


##### Insertion de graphiques (IMAGES PNG)
def add_chart_from_file(doc: Document, image_path: str, caption: str = None, width: float = 5.2):
    """
    Insère un graphique PNG dans le document, centré, avec légende optionnelle
    La largeur par défaut est ajustée pour éviter l'apparition de sauts de page forcés.
    Les graphiques sont générés par les modules app/charts et sauvegardés dans output/charts/
    """
    if not os.path.exists(image_path):
        # PLACEHOLDER si l'img est introuvable
        p = doc.add_paragraph()
        run = p.add_run(f"[Graphique non disponible : {os.path.basename(image_path)}]")
        run.italic = True
        run.font.color.rgb = RGBColor(150, 150, 150)
        p.paragraph_format.space_after = Pt(6)
        return # pour ne pas bloquer la génération du rapport
    
    # paragraphe centré pour l'image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    # insertion de l'img avec la largeur spécifiée (docx calcule la hauteur automatiquement pour garder les proportions)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width))

    # légende (optionnelle)
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)

        run_cap = p_cap.add_run(caption)
        run_cap.italic = True
        run_cap.font.size = Pt(9)
        run_cap.font.name = font_family
        run_cap.font.color.rgb = RGBColor(100, 100, 100)
        

##### Saut de page
def add_page_break(doc: Document):
    """
    Insère un saut de page dans le document: utilisé pour commencer chaque grande section sur une nouvelle page
    """
    doc.add_page_break()


##### Bloc de texte en cadré
def add_framed_text(doc: Document, text: str):
    """
    Ajoute un bloc de texte sur fond bleu cyan avec texte blanc.
    Pour des messages clés (exp: La temporalité des étudiants).
    Sous forme de: 1 cellule avec fond cyan, texte blanc à l'intérieur.
    """
    
    table = doc.add_table(rows=1, cols=1)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass # Le style n'existe pas dans le template, utiliser le style par défaut

    cell = table.cell(0, 0)
    cell.text = ""

    _set_cell_background(cell, cyan_hex)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(text)
    run.font.color.rgb = white
    run.font.size = Pt(11)
    run.font.name = font_family

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


##### Statistique clé
def add_key_stat(doc: Document, label: str, value: str):
    """
    Ajoute une statistique clé sous forme de: "Label: valeur"
    (exp: ► Total consultations : 5 230).
    """

    p = doc.add_paragraph()
    
    # symbole de puce 
    run_bullet = p.add_run("► ")
    run_bullet.font.color.rgb = cyan_primary
    run_bullet.font.name = font_family
    run_bullet.font.size = Pt(11)

    # label descriptif noir
    run_label = p.add_run(f"{label}: ")
    run_label.font.size = Pt(11)
    run_label.font.name = font_family
    run_label.font.color.rgb = black

    # valeur numérique en cyan gras
    run_value = p.add_run(value)
    run_value.font.bold = True
    run_value.font.color.rgb = cyan_primary
    run_value.font.name = font_family
    
    p.paragraph_format.space_after = Pt(4)


##### Sauvegarde du document
def save_document(doc: Document, output_path: str):
    """Sauvegarde le document word à l'emplacement spécifié."""
    # Crée le dossier parent si il n'existe pas encore
    parent_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(parent_dir, exist_ok=True)

    doc.save(output_path)


